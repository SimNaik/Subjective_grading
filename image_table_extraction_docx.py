from docx import Document
from docx.shared import Inches
import zipfile
import os
import re

# Function to remove content between < and >
def clean_text(text):
    # This regex will remove anything between < and >, including the brackets themselves
    return re.sub(r'<.*?>', '', text)

# Helper to get all block items (paragraphs and tables) in order
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

# Path to the input and output files
input_file_path = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/OCR_gd_gem/Converting Handwriting PDF to Word File/Biology Word file/06_10021024301039611141693746957.docx'
output_file_path = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/check_image/section_8.docx'
output_image_dir = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/check_image/'

doc = Document(input_file_path)
new_doc = Document()

# Function to extract images from DOCX
def extract_images_from_docx(docx_file_path, output_image_dir):
    with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
        image_dir = 'word/media/'
        image_files = [f for f in docx_zip.namelist() if f.startswith(image_dir)]
        if not image_files:
            print("No images found in the DOCX file.")
            return []
        extracted_images = []
        for image_file in image_files:
            image_data = docx_zip.read(image_file)
            image_filename = os.path.join(output_image_dir, os.path.basename(image_file))
            with open(image_filename, 'wb') as img_file:
                img_file.write(image_data)
            extracted_images.append(image_filename)
        return extracted_images

# Function to add image to new document
def add_image_to_doc(image_path, new_doc):
    new_doc.add_paragraph(f"Image: {os.path.basename(image_path)}")
    new_doc.add_picture(image_path, width=Inches(2))

# Function to copy table to new document
def copy_table(table, new_doc):
    new_table = new_doc.add_table(rows=0, cols=len(table.columns))
    for row in table.rows:
        new_row = new_table.add_row().cells
        for idx, cell in enumerate(row.cells):
            new_row[idx].text = cell.text

# Helper function to iterate through blocks (paragraphs and tables)
def iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# Extract images from the original DOCX file
extracted_images = extract_images_from_docx(input_file_path, output_image_dir)
image_index = 0

inside_sol = False
found_start = False

for block in iter_block_items(doc):
    if isinstance(block, Paragraph):
        text = block.text
        # Clean the text to remove content between <>
        cleaned_text = clean_text(text)

        if '<sol_start id=9>' in text and not found_start:
            inside_sol = True
            found_start = True
            continue
        if '<sol_end>' in text and inside_sol:
            inside_sol = False
            break
        if inside_sol:
            image_added = False
            for run in block.runs:
                if run._r.xml.find('graphic') != -1:
                    if image_index < len(extracted_images):
                        add_image_to_doc(extracted_images[image_index], new_doc)
                        image_index += 1
                    image_added = True
                    break
            if not image_added:
                new_doc.add_paragraph(cleaned_text)  # Add the cleaned text here
    elif isinstance(block, Table) and inside_sol:
        copy_table(block, new_doc)

# Save the new document with extracted content
new_doc.save(output_file_path)
print(f"Extracted content with images and tables has been saved to: {output_file_path}")
