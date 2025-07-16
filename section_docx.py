#This script extracts text and images from a DOCX file, then saves them into a new DOCX file. It extracts images, places them in the new document, and adds the corresponding text, including handling placeholders 
#for images. The final output is a DOCX file with both the text and images correctly inserted.
from docx import Document
from docx.shared import Inches
import zipfile
import os
import re

# Path to the input and output files
input_file_path = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/OCR_gd_gem/Converting Handwriting PDF to Word File/Biology Word file/06_10021024301039611141693746957.docx'
output_file_path = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/check_image/section_10.docx'
output_image_path = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/check_image/extracted_image.png'

# Load the input docx file
doc = Document(input_file_path)

# Create a new document to save the extracted content
new_doc = Document()

# Function to extract and save images from the DOCX file
def extract_images_from_docx(docx_file_path, output_image_dir):
    with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
        # Locate the image directory in the DOCX
        image_dir = 'word/media/'
        image_files = [f for f in docx_zip.namelist() if f.startswith(image_dir)]
        
        if not image_files:
            print("No images found in the DOCX file.")
            return []
        
        extracted_images = []
        for image_file in image_files:
            # Read the image data from the DOCX zip archive
            image_data = docx_zip.read(image_file)
            image_filename = os.path.join(output_image_dir, os.path.basename(image_file))
            
            # Save the image to the output directory
            with open(image_filename, 'wb') as img_file:
                img_file.write(image_data)
            
            extracted_images.append(image_filename)
        return extracted_images

# Function to add images to the new DOCX file at the correct place
def add_image_to_doc(image_path, new_doc):
    new_doc.add_paragraph(f"Image: {os.path.basename(image_path)}")  # Add image description
    new_doc.add_picture(image_path, width=Inches(2))  # Add image to new doc

# Function to remove unwanted tags (like <sol_start>, <sub_id>, etc.) from text
def clean_text(text):
    # Use regular expressions to remove all tags
    return re.sub(r'<.*?>', '', text)

# Extract images from the original DOCX file
extracted_images = extract_images_from_docx(input_file_path, '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/check_image/')

# Extract content and images together, using placeholders for images
inside_sol = False
image_index = 0  # Index to track which image to add next
content_with_placeholders = []

for paragraph in doc.paragraphs:
    if '<sol_start id=10>' in paragraph.text:
        inside_sol = True
        content_with_placeholders.append(clean_text(paragraph.text.replace('<sol_start id=10>', '').strip()))

    if inside_sol:
        # We need to capture the text before the image first
        image_added = False
        # Check if the paragraph contains an image
        for run in paragraph.runs:
            if run._r.xml.find('graphic') != -1:  # If an image is found
                # Add the text first and then image placeholder
                content_with_placeholders.append(clean_text(paragraph.text.strip()))  # Add text before the image
                content_with_placeholders.append("[image]")  # Placeholder for image
                image_added = True
                print(f"Image placeholder added after: {clean_text(paragraph.text.strip())}")  # Log for image
                break
        
        if not image_added:
            # Add the paragraph text (even if it's part of <sub_id>)
            content_with_placeholders.append(clean_text(paragraph.text.strip()))
            print(f"Text added: {clean_text(paragraph.text.strip())}")  # Log for text

    if '<sol_end>' in paragraph.text and inside_sol:
        inside_sol = False
        content_with_placeholders.append(clean_text(paragraph.text.replace('<sol_end>', '').strip()))
        print(f"End of section added: {clean_text(paragraph.text.strip())}")  # Log for end of section

# Add content to new DOCX
for text in content_with_placeholders:
    if text == "[image]" and image_index < len(extracted_images):
        # Add the corresponding image in place of the placeholder
        add_image_to_doc(extracted_images[image_index], new_doc)
        image_index += 1
        print(f"Image added at position {image_index}")  # Log for image insertion
    else:
        # Add the paragraph text (preserving everything, including text before image)
        new_doc.add_paragraph(text)
        print(f"Text inserted into new document: {text}")  # Log for text insertion

# Save the new document
new_doc.save(output_file_path)

print(f"Extracted content with images has been saved to: {output_file_path}")
