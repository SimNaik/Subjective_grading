from docx import Document
from docx.shared import Inches
import os
import zipfile
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image

input_base = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/OCR_gd_gem/final_evaluation/humans_ocr'
gemini_base = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/OCR_gd_gem/final_evaluation/gemini_768_ocr'
output_base = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/OCR_gd_gem/final_evaluation/tables'

def extract_images_from_docx(docx_file_path, output_image_dir):
    with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
        image_dir = 'word/media/'
        image_files = [f for f in docx_zip.namelist() if f.startswith(image_dir)]
        if not image_files:
            return []
        extracted_images = []
        for image_file in image_files:
            image_data = docx_zip.read(image_file)
            image_filename = os.path.join(output_image_dir, os.path.basename(image_file))
            with open(image_filename, 'wb') as img_file:
                img_file.write(image_data)
            try:
                with Image.open(image_filename) as img:
                    img.verify()
                extracted_images.append(image_filename)
            except Exception:
                os.remove(image_filename)
        return extracted_images

def copy_table(table, cell_doc):
    new_table = cell_doc.add_table(rows=0, cols=len(table.columns))
    for row in table.rows:
        new_row = new_table.add_row().cells
        for idx, cell in enumerate(row.cells):
            new_row[idx].text = cell.text

def iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def add_content_to_cell(doc_path, cell, image_dir):
    if not os.path.exists(doc_path):
        cell.add_paragraph("na")
        return
    doc = Document(doc_path)
    images = extract_images_from_docx(doc_path, image_dir)
    image_idx = 0
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            p = cell.add_paragraph(block.text)
            for run in block.runs:
                if 'graphic' in run._r.xml and image_idx < len(images):
                    try:
                        cell.add_paragraph().add_run().add_picture(images[image_idx], width=Inches(2))
                        image_idx += 1
                    except Exception as e:
                        print(f"Error adding image: {e}")
        elif isinstance(block, Table):
            copy_table(block, cell)

# Loop through all subfolders in humans_ocr
for folder_name in os.listdir(input_base):
    input_dir = os.path.join(input_base, folder_name)
    if not os.path.isdir(input_dir):
        continue
    output_dir = os.path.join(output_base, folder_name)
    os.makedirs(output_dir, exist_ok=True)
    section_files = sorted(
        [f for f in os.listdir(input_dir) if f.startswith('section_') and f.endswith('.docx') and not f.endswith('_table.docx')],
        key=lambda x: int(x.split('_')[1].split('.')[0])
    )
    for section_file in section_files:
        section_path = os.path.join(input_dir, section_file)
        gemini_section_path = os.path.join(gemini_base, folder_name, section_file)
        # Prepare image extraction dirs
        image_dir_human = os.path.join(output_dir, "images_human")
        image_dir_gemini = os.path.join(output_dir, "images_gemini")
        os.makedirs(image_dir_human, exist_ok=True)
        os.makedirs(image_dir_gemini, exist_ok=True)
        # Create a new document for this table
        table_doc = Document()
        table = table_doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.cell(0, 0).text = 'Human OCR'
        table.cell(0, 1).text = 'Gemini OCR'
        table.cell(0, 2).text = 'CER (Character Error Rate)'
        # Add all content to the first cell of the second row (Human OCR)
        cell_human = table.cell(1, 0)
        cell_human._element.clear_content()
        add_content_to_cell(section_path, cell_human, image_dir_human)
        # Add all content to the second cell of the second row (Gemini OCR)
        cell_gemini = table.cell(1, 1)
        cell_gemini._element.clear_content()
        add_content_to_cell(gemini_section_path, cell_gemini, image_dir_gemini)
        # Third cell: leave blank or set to 'na'
        table.cell(1, 2).text = ''
        # Save the table docx in the output subfolder with _table suffix
        table_filename = section_file.replace('.docx', '_table.docx')
        table_path = os.path.join(output_dir, table_filename)
        table_doc.save(table_path)
        print(f"Saved table for {section_file} to {table_path}")
