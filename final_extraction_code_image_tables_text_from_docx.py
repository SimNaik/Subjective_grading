from docx import Document
from docx.shared import Inches
import zipfile
import os
import re
from PIL import Image  # For image format verification
# Helper to get all block items (paragraphs and tables) in order
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

def clean_text(text):
    # Remove <sub_id = number> with optional spaces, line breaks, and extra spaces around the equal sign
    text = re.sub(r'<sub_id\s*=\s*\d+\s*>\s*', '', text)  # Match <sub_id = number> with spaces and line breaks

    # Remove <sub_id end> with optional spaces or line breaks around it
    text = re.sub(r'<sub_id\s*end\s*>\s*', '', text)  # Match <sub_id end> with spaces or line breaks

    return text


# Function to extract images from DOCX
def extract_images_from_docx(docx_file_path, output_image_dir):
    with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
        image_dir = 'word/media/'
        image_files = [f for f in docx_zip.namelist() if f.startswith(image_dir)]
        if not image_files:
            print(f"No images found in the DOCX file: {docx_file_path}")
            return []
        extracted_images = []
        for image_file in image_files:
            image_data = docx_zip.read(image_file)
            image_filename = os.path.join(output_image_dir, os.path.basename(image_file))
            with open(image_filename, 'wb') as img_file:
                img_file.write(image_data)
            # Check if the image is of a supported format
            try:
                # Attempt to open the image to verify it's a valid format
                with Image.open(image_filename) as img:
                    img.verify()  # Verify the image
                extracted_images.append(image_filename)
            except (IOError, SyntaxError) as e:
                print(f"Skipping unsupported or corrupted image: {image_filename}")
                os.remove(image_filename)  # Optionally delete the invalid image
        return extracted_images

# Function to add image to new document
def add_image_to_doc(image_path, new_doc):
    try:
        new_doc.add_paragraph(f"Image: {os.path.basename(image_path)}")
        new_doc.add_picture(image_path, width=Inches(2))
    except Exception as e:
        print(f"Error adding image {image_path}: {e}")

# Function to copy table to new document
def copy_table(table, new_doc):
    new_table = new_doc.add_table(rows=0, cols=len(table.columns))
    for row in table.rows:
        new_row = new_table.add_row().cells
        for idx, cell in enumerate(row.cells):
            new_row[idx].text = cell.text

# Helper function to iterate through blocks (paragraphs and tables)
def iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# Function to process sections for each DOCX file in the directory and save as separate files
def process_directory(input_dir, output_dir):
    # List all DOCX files in the input directory
    docx_files = [f for f in os.listdir(input_dir) if f.endswith('.docx')]
    
    # Process each DOCX file in the directory
    for docx_file in docx_files:
        input_file_path = os.path.join(input_dir, docx_file)
        print(f"Processing {input_file_path}")
        
        # Create a folder for each DOCX file using its base name
        base_name = os.path.splitext(docx_file)[0]
        output_section_dir = os.path.join(output_dir, base_name)
        os.makedirs(output_section_dir, exist_ok=True)
        
        # Create an images folder inside this directory
        output_image_dir = os.path.join(output_section_dir, "images")
        os.makedirs(output_image_dir, exist_ok=True)
        
        doc = Document(input_file_path)
        extracted_images = extract_images_from_docx(input_file_path, output_image_dir)
        image_index = 0
        inside_sol = False
        found_start = False
        section_id = None

        # Loop over all sections and create separate documents
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text
                cleaned_text = clean_text(text)  # Clean the text using the updated clean_text function

                # Find start of section based on ID tag
                if '<sol_start id=' in text and not found_start:
                    inside_sol = True
                    found_start = True
                    # Extract the section ID
                    section_id = re.search(r'<sol_start id=(\d+)>', text).group(1)
                    new_doc = Document()  # New document for each section
                    continue

                # End of section
                if '<sol_end>' in text and inside_sol:
                    inside_sol = False
                    if section_id:
                        # Save the document with section ID inside the folder created for the input file
                        output_file_path = os.path.join(output_section_dir, f"section_{section_id}.docx")
                        new_doc.save(output_file_path)
                        print(f"Extracted content for section {section_id} saved to: {output_file_path}")
                    found_start = False
                    continue

                # Add content inside section and remove <sub_id = number> and <sub_id end>
                if inside_sol:
                    # Clean the text inside the section
                    cleaned_text = re.sub(r'<sub_id\s*=\s*\d+\s*>\s*', '', cleaned_text)  # Remove <sub_id = number>
                    cleaned_text = re.sub(r'<sub_id\s*end\s*>\s*', '', cleaned_text)  # Remove <sub_id end>
                    
                    # Add the cleaned text to the new document
                    new_doc.add_paragraph(cleaned_text)  # Ensure cleaned text is added here

            elif isinstance(block, Table) and inside_sol:
                copy_table(block, new_doc)


# Function to extract and clean questions from docx
def extract_and_clean_questions(doc):
    question_text = []

    for para in doc.paragraphs:
        # Only keep the text between <sol_start> and <sol_end> tags (without the tags themselves)
        match = re.search(r'<sol_start.*?>\s*(.*?)\s*<sol_end>', para.text)
        if match:
            question_text_section = match.group(1)
            # Remove content between <sub_id = ...> and <sub_id end> tags but keep the text inside them
            question_text_cleaned = re.sub(r'<sub_id.*?>', '', question_text_section)  # Remove <sub_id ...>
            question_text_cleaned = re.sub(r'<sub_id end>', '', question_text_cleaned)  # Remove <sub_id end>
            question_text.append(question_text_cleaned.strip())  # Extract and clean question text
    
    return question_text


# Paths
input_dir = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/OCR_gd_gem/Converting Handwriting PDF to Word File/Biology Word file'
output_dir = '/Users/simrannaik/Desktop/solution_improvement/ds-prototypes/subjective_grading/solution_improvement/OCR_gd_gem/hujmanss_ocr'

# Process the directory containing DOCX files
process_directory(input_dir, output_dir)
